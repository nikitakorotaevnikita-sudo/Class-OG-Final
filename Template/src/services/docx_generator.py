"""Генерация Word-документа из markdown-ответа ИИ.

TEMPLATE: Адаптируй generate_report_docx() — замени заголовок документа и
поля info_rows под параметры проекта. Остальной код используй как есть.
"""

import io
import re
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


_COLOR_HEADER = "D9D9D9"  # серый — заголовки таблицы


def _set_cell_bg(cell, hex_color: str) -> None:
    """Устанавливает цвет фона ячейки таблицы."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tc_pr.append(shd)


def _set_row_bg(row, hex_color: str) -> None:
    """Устанавливает цвет фона для всех ячеек строки."""
    for cell in row.cells:
        _set_cell_bg(cell, hex_color)


def _cell_text(cell, text: str, bold: bool = False, size_pt: int = 10) -> None:
    """Устанавливает текст ячейки с поддержкой inline-форматирования."""
    cell.text = ""
    para = cell.paragraphs[0]
    _add_inline_runs(para, str(text))
    for run in para.runs:
        run.font.size = Pt(size_pt)
        if bold:
            run.bold = True


def _add_info_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    """Добавляет двухколоночную информационную таблицу без видимых границ."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "none")
                tc_borders.append(border)
            tc_pr.append(tc_borders)

    for i, (label, value) in enumerate(rows):
        label_cell = table.rows[i].cells[0]
        value_cell = table.rows[i].cells[1]
        _cell_text(label_cell, label + ":", bold=True, size_pt=10)
        _cell_text(value_cell, value, size_pt=10)
        label_cell.width = Cm(5)


def _is_table_row(line: str) -> bool:
    """Проверяет, является ли строка строкой markdown-таблицы."""
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and len(s) > 2


def _is_separator_row(line: str) -> bool:
    """Проверяет, является ли строка разделителем заголовка таблицы."""
    s = line.strip()
    if not (s.startswith("|") and s.endswith("|")):
        return False
    inner = s[1:-1]
    return all(re.match(r"^[:\-\s]+$", cell) for cell in inner.split("|"))


def _parse_table_row(line: str) -> list[str]:
    """Разбивает строку markdown-таблицы на ячейки."""
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [cell.strip() for cell in s.split("|")]


def _render_md_table(doc: Document, table_lines: list[str]) -> None:
    """Рендерит markdown-таблицу как Word-таблицу."""
    rows = [ln for ln in table_lines if not _is_separator_row(ln)]
    if not rows:
        return

    parsed = [_parse_table_row(r) for r in rows]
    col_count = max(len(r) for r in parsed)
    parsed = [r + [""] * (col_count - len(r)) for r in parsed]

    tbl = doc.add_table(rows=len(parsed), cols=col_count)
    tbl.style = "Table Grid"

    for row_idx, cells in enumerate(parsed):
        row = tbl.rows[row_idx]
        if row_idx == 0:
            _set_row_bg(row, _COLOR_HEADER)
        for col_idx, cell_text in enumerate(cells):
            cell = row.cells[col_idx]
            _cell_text(cell, cell_text, bold=(row_idx == 0), size_pt=9)

    doc.add_paragraph()


def _parse_markdown_to_doc(doc: Document, text: str) -> None:
    """Конвертирует markdown-текст в параграфы Word-документа."""
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        if _is_table_row(line):
            table_lines = []
            while i < len(lines) and (_is_table_row(lines[i]) or _is_separator_row(lines[i])):
                table_lines.append(lines[i])
                i += 1
            _render_md_table(doc, table_lines)
            continue

        if line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
            if p.runs:
                p.runs[0].font.size = Pt(11)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif re.match(r"^---+$", line.strip()):
            doc.add_paragraph("─" * 60)
        elif re.match(r"^[-*] ", line):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, line[2:].strip())
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, re.sub(r"^\d+\. ", "", line).strip())
        elif line.strip() == "":
            pass
        else:
            p = doc.add_paragraph()
            _add_inline_runs(p, line)

        i += 1


def _add_hyperlink(para, url: str, text: str) -> None:
    """Добавляет гиперссылку в параграф."""
    try:
        r_id = para.part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
    except Exception:
        para.add_run(text)
        return

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run_el = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    run_el.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    run_el.append(t)
    hyperlink.append(run_el)
    para._p.append(hyperlink)


def _add_inline_runs(para, text: str) -> None:
    """Добавляет текст с поддержкой **жирного**, *курсива* и [ссылок](url)."""
    pattern = re.compile(
        r"(\[([^\]]+)\]\(([^)]+)\)"
        r"|\*\*\*(.+?)\*\*\*"
        r"|\*\*(.+?)\*\*"
        r"|\*(.+?)\*"
        r"|`(.+?)`)"
    )
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            para.add_run(text[last:m.start()])
        raw = m.group(0)
        if raw.startswith("["):
            _add_hyperlink(para, m.group(3), m.group(2))
        elif raw.startswith("***"):
            run = para.add_run(m.group(4))
            run.bold = True
            run.italic = True
        elif raw.startswith("**"):
            run = para.add_run(m.group(5))
            run.bold = True
        elif raw.startswith("*"):
            run = para.add_run(m.group(6))
            run.italic = True
        else:
            run = para.add_run(m.group(7))
            run.font.name = "Courier New"
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])


def generate_report_docx(
    response_text: str,
    title: str = "Отчёт",
    **extra_info: str,
) -> bytes:
    """Генерирует Word-документ из markdown-ответа ИИ.

    Args:
        response_text: Текст отчёта в markdown-формате.
        title: Заголовок отчёта (название объекта, кейса и т.п.).
        **extra_info: Дополнительные поля для информационной таблицы
            (передавай как ключ-значение строки).

    Returns:
        bytes: Содержимое .docx файла.

    TEMPLATE: Адаптируй заголовок и поля info_rows под нужды проекта.
    Пример вызова с дополнительными полями:
        generate_report_docx(text, title="Проект X", period="Q1 2025")
    """
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    # TEMPLATE: Замени заголовок документа на релевантный для проекта
    title_para = doc.add_heading("Аналитический отчёт", level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    generated_at = datetime.now().strftime("%d.%m.%Y %H:%M")
    # TEMPLATE: Адаптируй info_rows — добавь/удали поля под параметры проекта
    info_rows: list[tuple[str, str]] = [("Объект", title)]
    for key, value in extra_info.items():
        info_rows.append((key, str(value)))
    info_rows.append(("Дата формирования", generated_at))

    _add_info_table(doc, info_rows)
    doc.add_paragraph()

    hr = doc.add_paragraph()
    hr_run = hr.add_run("─" * 80)
    hr_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    doc.add_paragraph()

    _parse_markdown_to_doc(doc, response_text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
